import os
import io
import pandas as pd
import fitz  # PyMuPDF
import docx
import openpyxl
from dotenv import load_dotenv

load_dotenv()

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 500))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 50))


def load_document(file_path: str) -> list[dict]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        pages = _load_pdf(file_path)
    elif ext == ".csv":
        pages = _load_csv(file_path)
    elif ext == ".txt":
        pages = _load_txt(file_path)
    elif ext == ".docx":
        pages = _load_docx(file_path)
    elif ext == ".xlsx":
        pages = _load_xlsx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    source = os.path.basename(file_path)
    chunks = []
    for page_text, page_num in pages:
        for chunk_text in _split_text(page_text):
            if chunk_text.strip():
                chunks.append({"text": chunk_text.strip(), "source": source, "page": page_num})
    return chunks


def _split_text(text: str) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        chunks.append(text[start:start + CHUNK_SIZE])
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def _load_pdf(file_path: str) -> list[tuple[str, int]]:
    doc = fitz.open(file_path)
    return [(page.get_text(), i + 1) for i, page in enumerate(doc)]


def _load_csv(file_path: str) -> list[tuple[str, int]]:
    df = pd.read_csv(file_path)
    rows = []
    for i, row in df.iterrows():
        text = ", ".join(f"{col}: {val}" for col, val in row.items())
        rows.append((text, int(i) + 1))
    return rows


def _load_txt(file_path: str) -> list[tuple[str, int]]:
    with open(file_path, encoding="utf-8", errors="replace") as f:
        return [(f.read(), 1)]


def _load_docx(file_path: str) -> list[tuple[str, int]]:
    doc = docx.Document(file_path)
    results = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            results.append((para.text, i + 1))
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            text = ", ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                results.append((text, i + 1))
    return results


def _load_xlsx(file_path: str) -> list[tuple[str, int]]:
    wb = openpyxl.load_workbook(file_path, data_only=True)
    results = []
    for sheet in wb.worksheets:
        headers = [cell.value for cell in next(sheet.iter_rows(min_row=1, max_row=1))]
        for i, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
            if not any(row):
                continue
            text = ", ".join(
                f"{h}: {v}" for h, v in zip(headers, row) if h is not None and v is not None
            )
            if text:
                results.append((text, f"{sheet.title}_{i}"))
    return results
